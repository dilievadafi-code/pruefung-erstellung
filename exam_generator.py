# -*- coding: utf-8 -*-
from __future__ import annotations
import json
import random
import tempfile
import zipfile
from pathlib import Path
from collections import Counter, defaultdict
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

LETTERS = ['A', 'B', 'C', 'D']
EXAM_TOPIC_REQUIREMENTS = {1:2, 2:2, 3:2, 4:3, 5:2, 6:8, 7:3, 8:5, 9:5, 10:5, 11:3, 12:2, 13:3, 14:3, 15:2}

class PoolError(Exception):
    pass

def load_pool(path: str | Path):
    data = json.loads(Path(path).read_text(encoding='utf-8'))
    validate_pool(data)
    return data

def validate_pool(data):
    questions = data.get('questions', [])
    if len(questions) < 50:
        raise PoolError('Der Fragenpool enthält zu wenige Fragen.')
    ids = [q.get('id') for q in questions]
    if len(ids) != len(set(ids)):
        raise PoolError('Der Fragenpool enthält doppelte IDs.')
    for q in questions:
        if q.get('topic_nr') not in EXAM_TOPIC_REQUIREMENTS:
            raise PoolError(f"Unbekanntes Themenfeld bei {q.get('id')}")
        if q.get('level') not in ('Basis', 'Experte'):
            raise PoolError(f"Ungültiger Schwierigkeitsgrad bei {q.get('id')}")
        if set(q.get('answers', {}).keys()) != set(LETTERS):
            raise PoolError(f"{q.get('id')} hat nicht genau vier Antworten A-D.")
        sol = q.get('solution', [])
        if not (1 <= len(sol) <= 3) or not set(sol) <= set(LETTERS):
            raise PoolError(f"{q.get('id')} hat keinen gültigen Lösungsschlüssel.")
    for topic_nr, needed in EXAM_TOPIC_REQUIREMENTS.items():
        if len([q for q in questions if q['topic_nr'] == topic_nr]) < needed:
            raise PoolError(f"Themenfeld {topic_nr}: zu wenige Fragen.")
    return True

def _expert_targets(count: int, rng: random.Random):
    # 25 Prozent von 50 Fragen sind 12,5. Einzelne Varianten nutzen daher 12 oder 13 Expertenfragen.
    targets = [12] * count
    total_experts = round(12.5 * count)
    extra = max(0, total_experts - 12 * count)
    positions = list(range(count))
    rng.shuffle(positions)
    for pos in positions[:extra]:
        targets[pos] = 13
    return targets

def _allocation(by_topic: dict, target_experts: int, rng: random.Random):
    topics = list(EXAM_TOPIC_REQUIREMENTS.keys())
    ranges = []
    for t in topics:
        needed = EXAM_TOPIC_REQUIREMENTS[t]
        basis = sum(1 for q in by_topic[t] if q['level'] == 'Basis')
        experte = sum(1 for q in by_topic[t] if q['level'] == 'Experte')
        lo = max(0, needed - basis)
        hi = min(needed, experte)
        if lo > hi:
            raise PoolError(f'Themenfeld {t}: Auswahl nach Niveau nicht möglich.')
        ranges.append((lo, hi))
    suffix_min = [0] * (len(topics) + 1)
    suffix_max = [0] * (len(topics) + 1)
    for i in range(len(topics) - 1, -1, -1):
        suffix_min[i] = suffix_min[i + 1] + ranges[i][0]
        suffix_max[i] = suffix_max[i + 1] + ranges[i][1]
    result = {}
    def rec(i, remaining):
        if i == len(topics):
            return remaining == 0
        t = topics[i]
        lo, hi = ranges[i]
        vals = list(range(lo, hi + 1))
        rng.shuffle(vals)
        for v in vals:
            if suffix_min[i + 1] <= remaining - v <= suffix_max[i + 1]:
                result[t] = v
                if rec(i + 1, remaining - v):
                    return True
        return False
    if not rec(0, target_experts):
        raise PoolError(f'Keine Auswahl mit {target_experts} Expertenfragen möglich.')
    return result

def select_questions(pool, rng: random.Random, exclude_ids=None, target_experts=13):
    exclude_ids = set(exclude_ids or [])
    by_topic = defaultdict(list)
    for q in pool['questions']:
        if q['id'] not in exclude_ids:
            by_topic[q['topic_nr']].append(q)
    # fallback: if exclusions make a topic infeasible, ignore exclusions for this attempt.
    try:
        allocation = _allocation(by_topic, target_experts, rng)
    except PoolError:
        by_topic = defaultdict(list)
        for q in pool['questions']:
            by_topic[q['topic_nr']].append(q)
        allocation = _allocation(by_topic, target_experts, rng)
    selected = []
    for t, needed in EXAM_TOPIC_REQUIREMENTS.items():
        e_needed = allocation[t]
        b_needed = needed - e_needed
        experts = [q for q in by_topic[t] if q['level'] == 'Experte']
        basis = [q for q in by_topic[t] if q['level'] == 'Basis']
        selected.extend(rng.sample(experts, e_needed))
        selected.extend(rng.sample(basis, b_needed))
    rng.shuffle(selected)
    return selected

def randomize_answers(q, rng: random.Random):
    order = LETTERS[:]
    rng.shuffle(order)
    new_answers = {}
    new_solution = []
    for idx, old in enumerate(order):
        new = LETTERS[idx]
        new_answers[new] = q['answers'][old]
        if old in q['solution']:
            new_solution.append(new)
    q2 = dict(q)
    q2['answers'] = new_answers
    q2['solution'] = sorted(new_solution)
    return q2

def _set_normal_style(doc):
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)

def _shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def create_exam_docx(questions, output_path, title, include_key=True, seed=None):
    rng = random.Random(seed)
    randomized = [randomize_answers(q, rng) for q in questions]
    rng.shuffle(randomized)
    doc = Document()
    _set_normal_style(doc)
    sec = doc.sections[0]
    sec.top_margin = Cm(1.7)
    sec.bottom_margin = Cm(1.6)
    sec.left_margin = Cm(1.9)
    sec.right_margin = Cm(1.9)
    footer = sec.footer.paragraphs[0]
    footer.text = '60 Minuten - 50 Multiple-Choice-Fragen - keine Hilfsmittel'
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(90,90,90)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(17)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('Dauer: 60 Minuten | Fragen: 50 | Bestehen ab 30 richtigen Fragen').italic = True
    doc.add_paragraph('Name: ___________________________________________    Datum: ____________________')
    doc.add_paragraph('Hinweis: Jede Frage hat vier Antwortmöglichkeiten. Eine, zwei oder drei Antworten können richtig sein. Eine Frage zählt nur, wenn alle richtigen und keine falschen Antworten markiert wurden.')
    if seed is not None:
        sp = doc.add_paragraph(f'Variante/Seed: {seed}')
        sp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in sp.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(100,100,100)
    doc.add_paragraph()
    answer_key = []
    for idx, q in enumerate(randomized, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(1)
        p.add_run(f'{idx}. Frage').bold = True
        qt = doc.add_paragraph(q['question'])
        qt.paragraph_format.space_after = Pt(2)
        for letter in LETTERS:
            op = doc.add_paragraph()
            op.paragraph_format.left_indent = Cm(0.6)
            op.paragraph_format.space_after = Pt(0)
            op.add_run(f'{letter.lower()}) ').bold = True
            op.add_run(q['answers'][letter])
        answer_key.append({'nr': idx, 'id': q['id'], 'topic_nr': q['topic_nr'], 'level': q['level'], 'solution': [x.lower() for x in q['solution']]})
    if include_key:
        doc.add_page_break()
        doc.add_heading('Lösungsschlüssel', level=1)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(['Frage', 'Lösung', 'ID', 'Thema', 'Niveau']):
            cell = table.rows[0].cells[i]
            cell.text = h
            _shade(cell, 'D9EAF7')
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for item in answer_key:
            cells = table.add_row().cells
            cells[0].text = str(item['nr'])
            cells[1].text = ', '.join(s + ')' for s in item['solution'])
            cells[2].text = item['id']
            cells[3].text = str(item['topic_nr'])
            cells[4].text = item['level']
            for cell in cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8)
    doc.save(output_path)
    return answer_key

def generate_exams(pool, count=1, title_prefix='IT-Grundschutz-Praktiker Übungsprüfung', seed=None, include_key=True, avoid_overlap=True, out_dir=None):
    rng = random.Random(seed)
    out_dir = Path(out_dir or tempfile.mkdtemp(prefix='itgs_exams_'))
    out_dir.mkdir(parents=True, exist_ok=True)
    paths = []
    logs = []
    targets = _expert_targets(count, rng)
    used = set()
    for i in range(1, count + 1):
        target = targets[i-1]
        exclude = used if avoid_overlap else set()
        selected = select_questions(pool, rng, exclude_ids=exclude, target_experts=target)
        if avoid_overlap:
            overlap = len(set(q['id'] for q in selected) & used)
            if overlap:
                logs.append(f'Variante {i}: {overlap} Fragen mussten wiederverwendet werden, weil der gewünschte Ausschluss nicht vollständig möglich war.')
            used.update(q['id'] for q in selected)
        variant_seed = rng.randint(1, 10**9)
        path = out_dir / f'pruefung_itgs_{i:02d}.docx'
        create_exam_docx(selected, path, f'{title_prefix} {i}', include_key=include_key, seed=variant_seed)
        levels = Counter(q['level'] for q in selected)
        topics = Counter(q['topic_nr'] for q in selected)
        logs.append(f"Variante {i}: {len(selected)} Fragen, {levels.get('Basis',0)} Basis, {levels.get('Experte',0)} Experte, Themenverteilung {dict(sorted(topics.items()))}.")
        paths.append(path)
    return paths, logs

def zip_files(paths, zip_path):
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as z:
        for p in paths:
            z.write(p, arcname=Path(p).name)
    return zip_path
